"""Generate small, fictional fixtures for Knowledge Agent dry-run tests."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
FIXTURES = ROOT / "knowledge" / "fixtures"
EVALS = ROOT / "knowledge" / "evals" / "questions.jsonl"


def write_image(path: Path) -> None:
    image = Image.new("RGB", (640, 240), "#dbeafe")
    draw = ImageDraw.Draw(image)
    draw.rectangle((24, 24, 616, 216), outline="#2563eb", width=4)
    draw.text((60, 100), "Orbit Knowledge Fixture Diagram", fill="#1e3a5f")
    image.save(path)


def write_markdown(path: Path) -> None:
    path.write_text(
        "# Starbridge Operations Policy\n\n"
        "## Support levels\n\n"
        "P1 incidents receive a first response within four hours.\n\n"
        "## Deletion policy\n\n"
        "Version 2026-Q3 supersedes Q2: deleted knowledge files are recoverable for 30 days.\n\n"
        "| Role | Can export |\n| --- | --- |\n| Owner | Yes |\n| Viewer | No |\n",
        encoding="utf-8",
    )


def write_docx(path: Path, image_path: Path, clean: bool) -> None:
    document = Document()
    if clean:
        document.add_heading("Starbridge Delivery Handbook", level=1)
        document.add_paragraph("This handbook describes the fictional Enterprise onboarding workflow.")
        document.add_heading("Milestones", level=2)
        table = document.add_table(rows=1, cols=3)
        for cell, value in zip(table.rows[0].cells, ["Stage", "Owner", "Exit condition"]):
            cell.text = value
        for row in [("Discovery", "Consultant", "SSO metadata received"), ("Launch", "PM", "Training completed")]:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = value
        document.add_picture(str(image_path), width=Inches(4.5))
        document.add_paragraph("Figure 1. A fictional delivery workflow illustration.")
    else:
        document.add_paragraph("notes from different meetings")
        document.add_paragraph("")
        document.add_paragraph("SSO metadata pending; maybe launch next month")
        document.add_paragraph("   repeated note: SSO metadata pending")
        document.add_paragraph("Owner said 30-day recovery? verify against Q3 policy.")
        document.add_paragraph("")
        document.add_picture(str(image_path), width=Inches(2.5))
    document.save(path)


def write_pdf(path: Path, image_path: Path, image_only: bool) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    if image_only:
        pdf.drawImage(str(image_path), 72, height - 310, width=460, height=172)
    else:
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(72, height - 72, "Starbridge Text Report")
        pdf.setFont("Helvetica", 11)
        pdf.drawString(72, height - 108, "Project ORB-2407 is blocked until SSO metadata is delivered.")
        pdf.drawString(72, height - 126, "Escalate to P1 if the metadata is absent after 2026-08-08.")
        pdf.showPage()
        pdf.setFont("Helvetica", 11)
        pdf.drawString(72, height - 72, "Page two: P1 first-response objective is four hours.")
    pdf.save()


def write_workbook(path: Path, clean: bool) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Projects" if clean else "Operations Notes"
    if clean:
        sheet.append(["Project ID", "Customer", "Risk", "Target date", "Description"])
        for cell in sheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D5E8F0")
        sheet.append(["ORB-2407", "Northstar", "High", "2026-08-15", "SSO metadata is pending"])
        sheet.append(["ORB-2411", "Seabank", "Medium", "2026-09-02", "Requirements confirmation"])
        sheet.freeze_panes = "A2"
    else:
        sheet.merge_cells("A1:D1")
        sheet["A1"] = "Operations notes - mixed layout"
        sheet["A1"].alignment = Alignment(horizontal="center")
        sheet.append([])
        sheet.append(["Team", "Issue", None, "Owner"])
        sheet.append(["Support", "P1 response slipped", None, "Lin"])
        sheet.append([])
        sheet.append(["Legacy", "Seven-day policy is obsolete", None, None])
        workbook.create_sheet("Scratchpad").append(["free form", "follow up"])
    workbook.save(path)


def write_evals() -> None:
    rows = [
        {"question": "What is the P1 first response target?", "expected": "four hours", "source": "clean-policy.md", "locator": "Support levels"},
        {"question": "How long is the Q3 deletion recovery period?", "expected": "30 days", "source": "clean-policy.md", "locator": "Deletion policy"},
        {"question": "What blocks ORB-2407?", "expected": "SSO metadata", "source": "text-report.pdf", "locator": "page 1"},
        {"question": "Which project has high risk?", "expected": "ORB-2407", "source": "clean-projects.xlsx", "locator": "Projects row 2"},
        {"question": "Can the scanned notice be indexed automatically?", "expected": "requires review", "source": "scanned-notice.pdf", "locator": "page 1"},
    ]
    EVALS.parent.mkdir(parents=True, exist_ok=True)
    EVALS.write_text("\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + "\n", encoding="utf-8")


def main() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    image_path = FIXTURES / "fixture-diagram.png"
    write_image(image_path)
    write_markdown(FIXTURES / "clean-policy.md")
    write_docx(FIXTURES / "clean-handbook.docx", image_path, clean=True)
    write_docx(FIXTURES / "messy-notes.docx", image_path, clean=False)
    write_pdf(FIXTURES / "text-report.pdf", image_path, image_only=False)
    write_pdf(FIXTURES / "scanned-notice.pdf", image_path, image_only=True)
    write_workbook(FIXTURES / "clean-projects.xlsx", clean=True)
    write_workbook(FIXTURES / "messy-operations.xlsx", clean=False)
    image_path.unlink()
    write_evals()


if __name__ == "__main__":
    main()
