"""Bounded, in-memory evidence extraction for Knowledge Agent prompts."""

from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        parts.extend(
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
        )
    return "\n".join(parts)


def _read_xlsx(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for worksheet in workbook.worksheets:
        parts.append(f"Worksheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def read_evidence(path: Path, *, max_chars: int = 6000) -> str:
    """Read a bounded content sample without persisting extracted content."""

    if max_chars < 0:
        raise ValueError("max_chars must be non-negative")
    readers = {
        ".md": _read_text,
        ".markdown": _read_text,
        ".txt": _read_text,
        ".docx": _read_docx,
        ".xlsx": _read_xlsx,
        ".pdf": _read_pdf,
    }
    reader = readers.get(path.suffix.lower())
    if reader is None:
        return ""
    return reader(path)[:max_chars]
