"""Deterministic local file profiling for Knowledge Agent planning."""

from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader

from .models import CorpusProfile


_EXTENSION_TYPES = {
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
}


def _relative_path(path: Path, root: Path) -> str:
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError as exc:
        raise ValueError("File is outside the requested knowledge folder") from exc


def _base_profile(path: Path, root: Path, file_type: str) -> dict[str, object]:
    content = path.read_bytes()
    return {
        "source_path": _relative_path(path, root),
        "source_hash": hashlib.sha256(content).hexdigest(),
        "file_type": file_type,
    }


def _profile_pdf(path: Path, values: dict[str, object]) -> CorpusProfile:
    reader = PdfReader(str(path))
    extracted = "".join(page.extract_text() or "" for page in reader.pages).strip()
    page_count = len(reader.pages)
    # A short text PDF still counts as usable; a scanned PDF normally extracts no text.
    ratio = min(1.0, len(extracted) / max(1, page_count * 200))
    return CorpusProfile(
        **values,
        page_count=page_count,
        text_extraction_ratio=ratio,
        image_count=sum(len(page.images) for page in reader.pages),
    )


def _profile_docx(path: Path, values: dict[str, object]) -> CorpusProfile:
    document = Document(str(path))
    heading_count = sum(
        1 for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")
    )
    image_count = len(document.inline_shapes)
    return CorpusProfile(
        **values,
        heading_count=heading_count,
        table_count=len(document.tables),
        image_count=image_count,
    )


def _profile_xlsx(path: Path, values: dict[str, object]) -> CorpusProfile:
    workbook = load_workbook(path, read_only=False, data_only=True)
    merged_cell_count = sum(len(sheet.merged_cells.ranges) for sheet in workbook.worksheets)
    blank_row_count = sum(
        1
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        if all(cell.value is None for cell in row)
    )
    table_quality = "messy" if merged_cell_count or blank_row_count else "stable"
    return CorpusProfile(
        **values,
        sheet_count=len(workbook.worksheets),
        merged_cell_count=merged_cell_count,
        blank_row_count=blank_row_count,
        table_quality=table_quality,
    )


def _profile_text(path: Path, values: dict[str, object]) -> CorpusProfile:
    content = path.read_text(encoding="utf-8", errors="replace")
    return CorpusProfile(
        **values,
        heading_count=sum(1 for line in content.splitlines() if line.lstrip().startswith("#")),
    )


def profile_file(path: Path, *, root: Path) -> CorpusProfile:
    """Profile one supported file, refusing paths outside *root*."""

    path = path.resolve()
    root = root.resolve()
    _relative_path(path, root)
    if not path.is_file():
        raise ValueError(f"Knowledge source is not a file: {path}")

    file_type = _EXTENSION_TYPES.get(path.suffix.lower(), "unknown")
    values = _base_profile(path, root, file_type)
    if file_type == "pdf":
        return _profile_pdf(path, values)
    if file_type == "docx":
        return _profile_docx(path, values)
    if file_type == "xlsx":
        return _profile_xlsx(path, values)
    if file_type in {"markdown", "text"}:
        return _profile_text(path, values)
    return CorpusProfile(**values)


def scan_folder(root: Path) -> list[CorpusProfile]:
    """Recursively profile supported files in a folder in stable path order."""

    root = root.resolve()
    if not root.is_dir():
        raise ValueError(f"Knowledge folder does not exist: {root}")
    files = sorted(
        (path for path in root.rglob("*") if path.is_file() and path.suffix.lower() in _EXTENSION_TYPES),
        key=lambda path: path.relative_to(root).as_posix(),
    )
    return [profile_file(path, root=root) for path in files]
